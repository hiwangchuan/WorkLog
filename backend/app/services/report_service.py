from io import BytesIO

import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


def markdown_bytes(content: str) -> bytes:
    return content.encode("utf-8")


def excel_bytes(records: list[dict]) -> bytes:
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        pd.DataFrame(records).to_excel(writer, index=False, sheet_name="WorkLog")
    buffer.seek(0)
    return buffer.getvalue()


def word_bytes(title: str, content: str) -> bytes:
    doc = Document()
    doc.add_heading(title, level=1)
    for line in content.splitlines():
        doc.add_paragraph(line)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def pdf_bytes(title: str, content: str) -> bytes:
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    font = "Helvetica"
    try:
        pdfmetrics.registerFont(TTFont("NotoSansCJK", "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"))
        font = "NotoSansCJK"
    except Exception:
        pass
    c.setFont(font, 16)
    c.drawString(48, height - 56, title)
    c.setFont(font, 10)
    y = height - 88
    for line in content.splitlines():
        if y < 48:
            c.showPage()
            c.setFont(font, 10)
            y = height - 56
        c.drawString(48, y, line[:95])
        y -= 18
    c.save()
    buffer.seek(0)
    return buffer.getvalue()
